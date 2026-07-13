import os
import re
from typing import Dict, Any, List, Optional

# Graceful imports for PDF and DOCX parsing
try:
    # pyrefly: ignore [missing-import]
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    # pyrefly: ignore [missing-import]
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    # pyrefly: ignore [missing-import]
    import docx
except ImportError:
    docx = None

# A comprehensive list of common industry skills for parsing extraction fallback
COMMON_SKILLS = [
    # Programming Languages
    "python", "javascript", "typescript", "java", "c++", "c#", "ruby", "php", "go", "rust", "scala", "kotlin", "swift", "sql", "r", "html", "css",
    # Frameworks & Libraries
    "react", "angular", "vue", "next.js", "django", "fastapi", "flask", "spring", "express", "node.js", "jquery", "bootstrap", "tailwind", "pandas", "numpy", "scikit-learn", "tensorflow", "pytorch", "keras",
    # Tools & Technologies
    "git", "github", "docker", "kubernetes", "aws", "azure", "gcp", "linux", "jenkins", "ansible", "terraform", "ci/cd", "graphql", "rest api",
    # Databases
    "postgresql", "mysql", "mongodb", "sqlite", "redis", "elasticsearch", "oracle",
    # Methodology & Soft Skills
    "agile", "scrum", "project management", "communication", "leadership", "problem solving", "teamwork", "product management", "machine learning", "deep learning", "nlp", "computer vision", "data analysis", "data science",
    # UI/UX & Design
    "figma", "sketch", "photoshop", "illustrator", "ui/ux", "wireframing"
]

def clean_text(text: str) -> str:
    """Clean and normalize extracted text."""
    if not text:
        return ""
    # Remove non-printable/weird control characters first
    text = ''.join(c for c in text if c.isprintable() or c in '\n\r\t')
    # Normalize whitespaces second
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def extract_text_from_pdf(file_path: str) -> str:
    """Extract all text from a PDF file using pdfplumber or PyMuPDF."""
    text = ""
    
    # Try pdfplumber first
    if pdfplumber:
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            print(f"pdfplumber failed: {e}. Trying PyMuPDF...")
            
    # Try PyMuPDF if pdfplumber fails or is not available
    if not text.strip() and fitz:
        try:
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text() + "\n"
        except Exception as e:
            print(f"PyMuPDF failed: {e}")
            
    return text

def extract_text_from_docx(file_path: str) -> str:
    """Extract all text from a DOCX file using python-docx."""
    text = ""
    if docx:
        try:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                if para.text:
                    text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            print(f"python-docx failed: {e}")
    else:
        print("python-docx is not installed.")
    return text

def parse_email(text: str) -> Optional[str]:
    email_regex = r'[\w\.-]+@[\w\.-]+\.\w+'
    match = re.search(email_regex, text)
    return match.group(0) if match else None

def parse_phone(text: str) -> Optional[str]:
    # Matches various phone styles like +1-555-555-5555, (555) 555-5555, etc.
    phone_regex = r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    match = re.search(phone_regex, text)
    return match.group(0) if match else None

def parse_name(text: str, email: str = None, phone: str = None) -> Optional[str]:
    """Heuristic to extract name: look at the top 4 lines of the resume text."""
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    if not lines:
        return None
        
    for line in lines[:4]:
        # Skip if contains email or phone
        if email and email in line:
            continue
        if phone and phone in line:
            continue
        # Skip line if it starts with common section headers
        lower_line = line.lower()
        if any(hdr in lower_line for hdr in ["curriculum", "resume", "cv", "experience", "education", "profile"]):
            continue
        # Check if line looks like a name (2-3 words, mostly alphabetic)
        words = line.split()
        if 1 <= len(words) <= 4:
            clean_line = re.sub(r'[^a-zA-Z\s]', '', line)
            if len(clean_line.strip()) > 3:
                return clean_line.strip()
    return "Candidate Name"

def extract_skills(text: str) -> List[str]:
    """Extract skills from text by checking against a skills vocabulary."""
    skills_found = []
    text_lower = text.lower()
    
    # Use word boundary search for exact matches
    for skill in COMMON_SKILLS:
        # Avoid simple short matches matching inside words (e.g. 'c' matching 'cats')
        if len(skill) <= 3:
            pattern = r'\b' + re.escape(skill) + r'\b'
        else:
            pattern = re.escape(skill)
            
        if re.search(pattern, text_lower):
            skills_found.append(skill)
            
    # Capitalize matched skills properly based on COMMON_SKILLS format
    mapped_skills = []
    for s in skills_found:
        # map to neat casings
        neat = {
            "python": "Python", "javascript": "JavaScript", "typescript": "TypeScript",
            "java": "Java", "c++": "C++", "c#": "C#", "ruby": "Ruby", "php": "PHP",
            "go": "Go", "rust": "Rust", "scala": "Scala", "kotlin": "Kotlin",
            "swift": "Swift", "sql": "SQL", "r": "R", "html": "HTML", "css": "CSS",
            "react": "React", "angular": "Angular", "vue": "Vue.js", "next.js": "Next.js",
            "django": "Django", "fastapi": "FastAPI", "flask": "Flask", "spring": "Spring Boot",
            "express": "Express.js", "node.js": "Node.js", "jquery": "jQuery",
            "bootstrap": "Bootstrap", "tailwind": "Tailwind CSS", "pandas": "Pandas",
            "numpy": "NumPy", "scikit-learn": "scikit-learn", "tensorflow": "TensorFlow",
            "pytorch": "PyTorch", "keras": "Keras", "git": "Git", "github": "GitHub",
            "docker": "Docker", "kubernetes": "Kubernetes", "aws": "AWS", "azure": "Azure",
            "gcp": "GCP", "linux": "Linux", "jenkins": "Jenkins", "ansible": "Ansible",
            "terraform": "Terraform", "ci/cd": "CI/CD", "graphql": "GraphQL",
            "rest api": "REST API", "postgresql": "PostgreSQL", "mysql": "MySQL",
            "mongodb": "MongoDB", "sqlite": "SQLite", "redis": "Redis",
            "elasticsearch": "Elasticsearch", "oracle": "Oracle DB", "agile": "Agile",
            "scrum": "Scrum", "project management": "Project Management", "communication": "Communication",
            "leadership": "Leadership", "problem solving": "Problem Solving", "teamwork": "Teamwork",
            "product management": "Product Management", "machine learning": "Machine Learning",
            "deep learning": "Deep Learning", "nlp": "Natural Language Processing (NLP)",
            "computer vision": "Computer Vision", "data analysis": "Data Analysis",
            "data science": "Data Science", "figma": "Figma", "sketch": "Sketch",
            "photoshop": "Photoshop", "illustrator": "Illustrator", "ui/ux": "UI/UX",
            "wireframing": "Wireframing"
        }.get(s, s.title())
        mapped_skills.append(neat)
        
    return list(set(mapped_skills))

def extract_section(text: str, keywords: List[str]) -> str:
    """Extract text from a section based on heading keywords."""
    lines = text.split('\n')
    section_content = []
    in_section = False
    
    # Common headers that would signal the end of a section
    other_headers = ["experience", "education", "skills", "projects", "certifications", "languages", "summary", "profile", "contact", "about me", "work history", "employment"]
    
    for line in lines:
        cleaned_line = line.strip().lower()
        if not cleaned_line:
            continue
            
        # Check if this line is the heading we want
        if any(re.search(r'\b' + re.escape(kw) + r'\b', cleaned_line) for kw in keywords):
            in_section = True
            continue
            
        # Check if we hit another section heading
        if in_section:
            # If the line looks like a header and contains a header keyword that isn't in our current target keywords
            if any(re.search(r'\b' + re.escape(hdr) + r'\b', cleaned_line) for hdr in other_headers if hdr not in keywords):
                # Check if it is a short line (typical of headers)
                if len(cleaned_line) < 30:
                    in_section = False
                    
        if in_section:
            section_content.append(line)
            
    return "\n".join(section_content).strip()

def parse_resume(file_path: str) -> Dict[str, Any]:
    """Parse resume from PDF/DOCX and return structured components."""
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        raw_text = extract_text_from_pdf(file_path)
    elif ext == ".docx":
        raw_text = extract_text_from_docx(file_path)
    else:
        raw_text = ""
        
    cleaned_text = clean_text(raw_text)
    
    email = parse_email(cleaned_text)
    phone = parse_phone(cleaned_text)
    name = parse_name(raw_text, email, phone)
    skills = extract_skills(cleaned_text)
    
    # Extract sections
    education = extract_section(raw_text, ["education", "academic", "university", "college", "degree"])
    experience = extract_section(raw_text, ["experience", "employment", "work history", "job history", "career"])
    projects = extract_section(raw_text, ["projects", "personal projects", "academic projects"])
    certifications = extract_section(raw_text, ["certifications", "certificates", "credentials", "licenses"])
    
    # Extract languages
    languages_list = []
    languages_keywords = ["english", "spanish", "french", "german", "mandarin", "chinese", "hindi", "japanese", "arabic", "portuguese", "russian", "italian"]
    for lang in languages_keywords:
        if re.search(r'\b' + re.escape(lang) + r'\b', cleaned_text.lower()):
            languages_list.append(lang.title())
            
    # Estimate years of experience from experience text
    years_exp = 0
    # Try simple regex: "X years", "X+ years", "X years of experience"
    exp_matches = re.findall(r'(\d+)\+?\s*years?\s+(?:of\s+)?experience', cleaned_text.lower())
    if exp_matches:
        years_exp = max(int(m) for m in exp_matches)
    else:
        # Fallback heuristic: count number of jobs/ranges (e.g. 2018-2022)
        date_ranges = re.findall(r'\b(19\d{2}|20\d{2})\s*[-–—]\s*(20\d{2}|present)\b', cleaned_text.lower())
        total_years = 0
        for start, end in date_ranges:
            start_yr = int(start)
            end_yr = 2026 if end == "present" else int(end) # Using current local time year (2026)
            diff = end_yr - start_yr
            if 0 < diff < 15:
                total_years += diff
        if total_years > 0:
            years_exp = min(total_years, 20) # Cap at 20
            
    return {
        "name": name,
        "email": email,
        "phone": phone,
        "skills": ", ".join(skills) if skills else "",
        "education": education if education else "Self-taught / Other",
        "experience": experience if experience else f"{years_exp} years of experience" if years_exp else "Entry level",
        "projects": projects if projects else "",
        "certifications": certifications if certifications else "",
        "languages": ", ".join(languages_list) if languages_list else "English",
        "years_experience": years_exp
    }
